"""Report generation endpoints — PDF and DOCX."""
import io
from datetime import datetime

from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from config import SESSIONS_DIR, USER_FILE, TESTS_DIR, FONTS_DIR
from storage import read_json, file_exists

router = APIRouter()


class ReportRequest(BaseModel):
    session_id: str
    format: str  # "pdf" or "docx"


def _load_report_data(session_id: str):
    """Load and assemble all data needed for a report."""
    session_path = SESSIONS_DIR / f"{session_id}.json"
    if not file_exists(session_path):
        raise HTTPException(status_code=404, detail="Session not found")

    session = read_json(session_path)
    user = read_json(USER_FILE) if file_exists(USER_FILE) else {}

    # Find test
    test = None
    if TESTS_DIR.exists():
        for topic_dir in TESTS_DIR.iterdir():
            if not topic_dir.is_dir():
                continue
            test_file = topic_dir / f"{session['test_id']}.json"
            if file_exists(test_file):
                test = read_json(test_file)
                break

    if not test:
        raise HTTPException(status_code=404, detail="Test not found for this session")

    return session, user, test


def _build_pdf(session, user, test) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    )
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    # Register Cyrillic-capable fonts bundled with the app
    font_regular = str(FONTS_DIR / 'DejaVuSans.ttf')
    font_bold    = str(FONTS_DIR / 'DejaVuSans-Bold.ttf')
    pdfmetrics.registerFont(TTFont('DejaVu',     font_regular))
    pdfmetrics.registerFont(TTFont('DejaVu-Bold', font_bold))

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)

    base = {'fontName': 'DejaVu', 'fontSize': 10, 'spaceAfter': 4, 'leading': 14}
    title_style   = ParagraphStyle('RTitle',   fontName='DejaVu-Bold', fontSize=16, spaceAfter=6, leading=22)
    heading_style = ParagraphStyle('RHeading', fontName='DejaVu-Bold', fontSize=12, spaceAfter=4, leading=16)
    body_style    = ParagraphStyle('RBody',    **base)
    correct_style = ParagraphStyle('RCorrect', textColor=colors.HexColor('#2e7d32'), **base)
    wrong_style   = ParagraphStyle('RWrong',   textColor=colors.HexColor('#c62828'), **base)

    story = []

    # Header
    story.append(Paragraph("Тренажёр антикризисного управления", title_style))
    story.append(Paragraph("Отчёт о прохождении теста", heading_style))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.grey))
    story.append(Spacer(1, 0.3*cm))

    # Student info
    full_name = f"{user.get('last_name', '')} {user.get('first_name', '')}".strip()
    story.append(Paragraph(f"Студент: {full_name}", body_style))
    story.append(Paragraph(f"Группа: {user.get('group', '—')}", body_style))
    finished = session.get('finished_at', '')[:10]
    story.append(Paragraph(f"Дата: {finished}", body_style))
    story.append(Spacer(1, 0.2*cm))

    # Test info
    story.append(Paragraph(f"Тест: {test['title']}", heading_style))
    story.append(Spacer(1, 0.3*cm))

    # Build answer map
    answer_map = {a['question_id']: a for a in session.get('answers', [])}

    correct_count = 0
    total = len(test.get('questions', []))

    for q in test.get('questions', []):
        story.append(Paragraph(f"Вопрос {q['order']}: {q['text']}", heading_style))

        session_answer = answer_map.get(q['id'], {})
        selected_ids = set(session_answer.get('selected_answer_ids', []))

        correct_ids = {a['id'] for a in q['answers'] if a['is_correct']}
        is_fully_correct = selected_ids == correct_ids

        if is_fully_correct:
            correct_count += 1

        for ans in q['answers']:
            marker = ""
            if ans['is_correct'] and ans['id'] in selected_ids:
                marker = "✓ "
                style = correct_style
            elif ans['id'] in selected_ids and not ans['is_correct']:
                marker = "✗ "
                style = wrong_style
            elif ans['is_correct']:
                marker = "→ "
                style = correct_style
            else:
                style = body_style

            story.append(Paragraph(f"  {marker}{ans['text']}", style))

        comment = session_answer.get('comment', '')
        if comment:
            story.append(Paragraph(f"Комментарий студента: {comment}", body_style))

        explanation = q.get('explanation', '')
        if explanation:
            story.append(Paragraph(f"Пояснение: {explanation}", body_style))
        story.append(Spacer(1, 0.3*cm))

    # Summary
    story.append(HRFlowable(width="100%", thickness=1, color=colors.grey))
    story.append(Spacer(1, 0.2*cm))
    pct = round(correct_count / total * 100) if total else 0
    story.append(Paragraph(f"Результат: {correct_count} из {total} ({pct}%)", heading_style))
    story.append(Spacer(1, 0.2*cm))
    story.append(Paragraph(f"Что сделано верно: {test.get('analysis_good', '')}", body_style))
    story.append(Paragraph(f"Что стоит улучшить: {test.get('analysis_improve', '')}", body_style))

    doc.build(story)
    return buf.getvalue()


def _build_docx(session, user, test) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    doc.add_heading("Тренажёр антикризисного управления", 0)
    doc.add_heading("Отчёт о прохождении теста", 1)

    full_name = f"{user.get('last_name', '')} {user.get('first_name', '')}".strip()
    doc.add_paragraph(f"Студент: {full_name}")
    doc.add_paragraph(f"Группа: {user.get('group', '—')}")
    finished = session.get('finished_at', '')[:10]
    doc.add_paragraph(f"Дата: {finished}")
    doc.add_paragraph(f"Тест: {test['title']}")
    doc.add_paragraph("")

    answer_map = {a['question_id']: a for a in session.get('answers', [])}
    correct_count = 0
    total = len(test.get('questions', []))

    for q in test.get('questions', []):
        doc.add_heading(f"Вопрос {q['order']}: {q['text']}", 2)

        session_answer = answer_map.get(q['id'], {})
        selected_ids = set(session_answer.get('selected_answer_ids', []))
        correct_ids = {a['id'] for a in q['answers'] if a['is_correct']}
        is_fully_correct = selected_ids == correct_ids
        if is_fully_correct:
            correct_count += 1

        for ans in q['answers']:
            p = doc.add_paragraph()
            if ans['is_correct'] and ans['id'] in selected_ids:
                run = p.add_run(f"✓ {ans['text']}")
                run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
            elif ans['id'] in selected_ids and not ans['is_correct']:
                run = p.add_run(f"✗ {ans['text']}")
                run.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
            elif ans['is_correct']:
                run = p.add_run(f"→ {ans['text']}")
                run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
            else:
                p.add_run(f"  {ans['text']}")

        comment = session_answer.get('comment', '')
        if comment:
            doc.add_paragraph(f"Комментарий: {comment}")

        explanation = q.get('explanation', '')
        if explanation:
            doc.add_paragraph(f"Пояснение: {explanation}")

    doc.add_heading("Итог", 1)
    pct = round(correct_count / total * 100) if total else 0
    doc.add_paragraph(f"Результат: {correct_count} из {total} ({pct}%)")
    doc.add_paragraph(f"Что сделано верно: {test.get('analysis_good', '')}")
    doc.add_paragraph(f"Что стоит улучшить: {test.get('analysis_improve', '')}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@router.post("/generate")
def generate_report(req: ReportRequest):
    session, user, test = _load_report_data(req.session_id)

    if req.format == "pdf":
        data = _build_pdf(session, user, test)
        media_type = "application/pdf"
        filename = f"report_{req.session_id}.pdf"
    elif req.format == "docx":
        data = _build_docx(session, user, test)
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        filename = f"report_{req.session_id}.docx"
    else:
        raise HTTPException(status_code=400, detail="Format must be 'pdf' or 'docx'")

    return StreamingResponse(
        io.BytesIO(data),
        media_type=media_type,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )
